"""Ingestion tests: PDF, DOCX, text, garbage bytes."""

from app.resume.ingest import detect_kind, extract_raw_text


def test_text_precedence():
    assert extract_raw_text(data=b"ignored", text="provided text") == "provided text"


def test_empty_input():
    assert extract_raw_text() == ""
    assert extract_raw_text(data=b"") == ""


def test_detect_kind():
    assert detect_kind(b"%PDF-1.4", None) == "pdf"
    assert detect_kind(b"PK\x03\x04junk", None) == "docx"
    assert detect_kind(b"\x00\x01", None) == "unknown"
    assert detect_kind(b"", "application/pdf") == "pdf"


def test_docx_extraction():
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Built a React dashboard.")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_raw_text(data=buf.getvalue())
    assert "Jane Doe" in text
    assert "React" in text


def test_garbage_bytes_returns_empty():
    assert extract_raw_text(data=b"\x00\x01\x02 not a pdf") == ""


def test_long_text_is_truncated():
    text = extract_raw_text(text="x" * 50000)
    assert len(text) <= 20000
