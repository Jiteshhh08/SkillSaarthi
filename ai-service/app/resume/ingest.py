"""Resume ingestion — file bytes → raw resume text.

PDF text layers are extracted with pypdf. DOCX files are converted with
python-docx. Image-only/scanned PDFs yield an empty text layer — the LLM
extraction then reports an empty resume rather than hallucinating one.
"""

import io
import logging
import re
import zipfile

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Letter-spacing densifier (collapses per-glyph spaced PDF text layers)
# ---------------------------------------------------------------------------
_TOKEN_RE = r"\w\.@_"
_LETTER_SPACED_RE = re.compile(r"(?<=[" + _TOKEN_RE + r"]) (?=[" + _TOKEN_RE + r"])")
_SPACE_RUN_RE = re.compile(r"[ ]{2,}")
_ALNUM_RUN_RE = re.compile(r"[\w.@_]+")


def _is_letter_spaced(text):
    """Heuristic: are words rendered as isolated single characters?

    In a letter-spaced text layer, every glyph is followed by one space, so
    there are almost no multi-character runs. Normal prose, by contrast, is
    dominated by real words (multi-character runs between spaces).
    """
    if not text:
        return False
    runs = _ALNUM_RUN_RE.findall(text)
    if len(runs) < 8:
        return False
    multi_char = sum(1 for run in runs if len(run) >= 2)
    return multi_char / len(runs) < 0.05


def densify_text(text):
    """Collapse the per-glyph spaces of a letter-spaced PDF text layer.

    Words in such extractions appear as e.g. 'j a v a s c r i p t'. We remove a
    single space that sits between two word characters, leaving multi-space
    runs (which separate real words) and newlines intact.
    """
    if not text or not _is_letter_spaced(text):
        return text
    chunks = re.split(r"( {2,}|\n)", text)
    rebuilt = []
    for chunk in chunks:
        if chunk is None:
            continue
        if re.fullmatch(r"( {2,}|\n)", chunk):
            rebuilt.append(chunk)
        else:
            rebuilt.append(_LETTER_SPACED_RE.sub("", chunk))
    result = "".join(rebuilt)
    result = re.sub(r"[ ]{2,}", " ", result)
    result = re.sub(r"\s+([.,;:!?'\"\)\]])", r"\1", result)
    result = re.sub(r"([(\[])\s+", r"\1", result)
    return result

MAX_RAW_TEXT = 20000

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"


def detect_kind(data, mime_type):
    if mime_type:
        mime = mime_type.lower()
        if "pdf" in mime:
            return "pdf"
        if mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
            return "docx" if "openxml" in mime else "doc"
    if data[:4] == _PDF_MAGIC:
        return "pdf"
    if data[:4] == _ZIP_MAGIC:
        return "docx"
    return "unknown"


def _extract_pdf(data):
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — a single bad page must not kill extraction
            continue
    return densify_text("\n".join(pages))


def _extract_docx(data):
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_raw_text(data=None, text=None, mime_type=None):
    """Return up to MAX_RAW_TEXT chars of resume text.

    ``text`` (pre-extracted by the caller) takes precedence — this matches
    the existing analyzer contract.
    """
    if text and text.strip():
        return text[:MAX_RAW_TEXT]

    if not data:
        return ""

    if zipfile.is_zipfile(io.BytesIO(data)):
        try:
            return _extract_docx(data)[:MAX_RAW_TEXT]
        except Exception as error:  # noqa: BLE001
            logger.warning("DOCX parse failed: %s", error)
            return ""

    kind = detect_kind(data, mime_type)
    if kind == "pdf":
        try:
            return _extract_pdf(data)[:MAX_RAW_TEXT]
        except Exception as error:  # noqa: BLE001
            logger.warning("PDF parse failed: %s", error)
            return ""

    # Unknown or unsupported format.
    return ""